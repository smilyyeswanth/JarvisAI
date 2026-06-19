
from fastapi import FastAPI, UploadFile, File, Form
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from agents.report_generator import create_report

from agents.job_agent import search_jobs
from agents.resume_agent import analyze_job_description
from agents.resume_parser import read_resume
from agents.resume_matcher import compare_resume_to_jd
from ai.ollama_client import ask_ai
from actions.vision import analyze_screen
from actions.websites import *

from agents.resume_tailor import tailor_resume
from docx import Document
import os
import socket
import webbrowser
from datetime import datetime


# =====================================
# FASTAPI APP
# =====================================

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =====================================
# REQUEST MODEL
# =====================================

class Prompt(BaseModel):
    message: str


# =====================================
# APP COMMANDS
# =====================================

def open_chrome():
    os.system("open -a 'Google Chrome'")


def open_safari():
    os.system("open -a Safari")


def open_notes():
    os.system("open -a Notes")


def open_vscode():
    os.system("open -a 'Visual Studio Code'")


def open_finder():
    os.system("open -a Finder")


# =====================================
# SYSTEM COMMANDS
# =====================================

def get_time():
    return datetime.now().strftime("%I:%M %p")


def get_ip():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "Unable to determine IP"


def get_battery():
    try:
        result = os.popen("pmset -g batt").read()

        for line in result.split("\n"):
            if "%" in line:
                return line.strip()

        return "Battery information unavailable"

    except Exception:
        return "Battery information unavailable"




# =====================================
# FILE SEARCH
# =====================================

SEARCH_DIRS = [
    os.path.expanduser("~/Desktop"),
    os.path.expanduser("~/Documents"),
    os.path.expanduser("~/Downloads"),
]


def find_file(filename):
    matches = []

    for directory in SEARCH_DIRS:
        for root, dirs, files in os.walk(directory):
            for file in files:
                if filename.lower() in file.lower():
                    matches.append(os.path.join(root, file))

    return matches


def open_file(path):
    os.system(f'open "{path}"')

# =====================================
# TAILORED RESUME GENERATOR
# =====================================

def create_tailored_resume(content):

    doc = Document()

    doc.add_heading(
        "TAILORED RESUME",
        level=1
    )

    doc.add_paragraph(content)

    output_file = os.path.expanduser(
        "~/Desktop/Tailored_Resume.docx"
    )

    doc.save(output_file)

    return output_file

# =====================================
# ROUTES
# =====================================

@app.get("/")
def home():
    return {
        "message": "Jarvis Backend Online"
    }


@app.post("/resume-analysis")
def resume_analysis(prompt: Prompt):

    result = analyze_job_description(
        prompt.message
    )

    return {
        "response": result
    }


@app.post("/resume-match")
async def resume_match(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    try:
        upload_dir = "uploads"
        os.makedirs(
            upload_dir,
            exist_ok=True
        )

        file_path = os.path.join(
            upload_dir,
            file.filename
        )

        with open(
            file_path,
            "wb"
        ) as buffer:
            buffer.write(
                await file.read()
            )

        resume_text = read_resume(
            file_path
        )
        result = compare_resume_to_jd(
    resume_text,
    job_description
)

        report_path = create_report(
        result
        )

        return {
    "response": result,
    "report": report_path
}



    except Exception as e:
        return {
            "response": f"Resume Match Error: {str(e)}"
        }

@app.post("/tailor-resume")
async def tailor_resume_endpoint(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):

    try:

        upload_dir = "uploads"

        os.makedirs(
            upload_dir,
            exist_ok=True
        )

        file_path = os.path.join(
            upload_dir,
            file.filename
        )

        with open(
            file_path,
            "wb"
        ) as buffer:

            buffer.write(
                await file.read()
            )

        resume_text = read_resume(
            file_path
        )

        tailored_resume = tailor_resume(
            resume_text,
            job_description
        )

        output_file = create_tailored_resume(
            tailored_resume
        )

        return {
            "response":
            "Tailored resume created successfully",
            "file":
            output_file
        }

    except Exception as e:

        return {
            "response":
            f"Tailor Resume Error: {str(e)}"
        }

@app.post("/chat")
def chat(prompt: Prompt):

    try:
        user_text = (
            prompt.message.lower()
            .replace("jarvis", "")
            .strip()
        )

        print(f"User: {user_text}")

        # =================================
        # SCREEN ANALYSIS
        # =================================

        if (
            "what's on my screen" in user_text
            or "what is on my screen" in user_text
            or "analyze screen" in user_text
            or "what's on the screen" in user_text
            or "what is on the screen" in user_text
            or user_text == "screen"
        ):

            result = analyze_screen()

            return {
                "response": result
            }

        # =================================
        # APP COMMANDS
        # =================================

        if "open chrome" in user_text:
            open_chrome()
            return {"response": "Opening Chrome"}

        if "open safari" in user_text:
            open_safari()
            return {"response": "Opening Safari"}

        if "open notes" in user_text:
            open_notes()
            return {"response": "Opening Notes"}

        if (
            "open vscode" in user_text
            or "open visual studio code" in user_text
        ):
            open_vscode()
            return {"response": "Opening Visual Studio Code"}

        if "open finder" in user_text:
            open_finder()
            return {"response": "Opening Finder"}

        # =================================
        # WEBSITE COMMANDS
        # =================================

        if "open youtube" in user_text:
            open_youtube()
            return {"response": "Opening YouTube"}

        if "open linkedin" in user_text:
            open_linkedin()
            return {"response": "Opening LinkedIn"}

        if "open github" in user_text:
            open_github()
            return {"response": "Opening GitHub"}

        if "open gmail" in user_text:
            open_gmail()
            return {"response": "Opening Gmail"}

        if "open chatgpt" in user_text:
            open_chatgpt()
            return {"response": "Opening ChatGPT"}

        if "open google" in user_text:
            webbrowser.open("https://google.com")
            return {"response": "Opening Google"}

        # =================================
        # FILE CREATION
        # =================================

        if "create file" in user_text:

            filename = (
                user_text
                .replace("create file", "")
                .replace("called", "")
                .strip()
            )

            if not filename:
                filename = "new_file.txt"

            desktop = os.path.expanduser("~/Desktop")

            filepath = os.path.join(
                desktop,
                filename
            )

            with open(filepath, "w") as f:
                f.write("")

            return {
                "response": f"Created {filename} on your Desktop"
            }

        # =================================
        # JOB SEARCH
        # =================================

        if "find jobs" in user_text.lower():
            print("JOB SEARCH TRIGGERED")

            keyword = (
                user_text.lower()
                .replace("find jobs", "")
                .strip()
            )

            print("KEYWORD:", keyword)
            jobs = search_jobs(keyword)
            print("RESULT:", jobs)

            return {
                "response": str(jobs)
            }

        # =================================
        # SYSTEM INFO
        # =================================

        if (
            "what time" in user_text
            or "current time" in user_text
        ):
            return {
                "response": f"The current time is {get_time()}"
            }

        if "battery" in user_text:
            return {
                "response": f"Battery status: {get_battery()}"
            }

        if "ip address" in user_text:
            return {
                "response": f"Your IP address is {get_ip()}"
            }

        # =================================
        # SCREENSHOT
        # =================================

        if "take screenshot" in user_text:

            path = "/tmp/jarvis_screenshot.png"

            os.system(
                f"screencapture -x {path}"
            )

            return {
                "response": f"Screenshot saved to {path}"
            }

        # =================================
        # CLIPBOARD
        # =================================

        if "clipboard" in user_text:

            text = os.popen(
                "pbpaste"
            ).read()

            return {
                "response": f"Clipboard contains: {text}"
            }

        # =================================
        # FILE SEARCH
        # =================================

        if user_text.startswith("find"):

            search_text = (
                user_text
                .replace("find", "")
                .strip()
            )

            files = find_file(search_text)

            if files:

                open_file(files[0])

                return {
                    "response": f"Found and opened:\n{files[0]}"
                }

            return {
                "response": f"Could not find {search_text}"
            }

        # =================================
        # CREATE FOLDER
        # =================================

        if "create folder" in user_text:

            folder_name = (
                user_text
                .replace("create folder", "")
                .replace("called", "")
                .strip()
            )

            if not folder_name:
                folder_name = "New Folder"

            desktop = os.path.expanduser(
                "~/Desktop"
            )

            folder_path = os.path.join(
                desktop,
                folder_name
            )

            os.makedirs(
                folder_path,
                exist_ok=True
            )

            return {
                "response": f"Created folder {folder_name}"
            }       

        # =================================
        # AI FALLBACK
        # =================================

        response = ask_ai(prompt.message)

        return {
            "response": response
        }

    except Exception as e:

        print("ERROR:", str(e))

        return {
            "response": f"Jarvis Error: {str(e)}"
        }

