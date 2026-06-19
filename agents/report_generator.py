from docx import Document
import os

def create_report(content):

    doc = Document()

    doc.add_heading(
        "JARVIS ATS REPORT",
        level=1
    )

    doc.add_paragraph(content)

    desktop = os.path.expanduser(
        "~/Desktop"
    )

    output_file = os.path.join(
        desktop,
        "ATS_Report.docx"
    )

    doc.save(output_file)

    return output_file